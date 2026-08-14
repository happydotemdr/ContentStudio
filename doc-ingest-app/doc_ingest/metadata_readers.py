"""Independent (non-firecrawl) readers for the gauntlet's integrity checks
(spec §8) -- page/word/table/row counts computed from the source file
directly, never re-derived from firecrawl-parse's own output, so the check
is independent of the thing it's checking."""
from __future__ import annotations

from pathlib import Path

import docx
import openpyxl
from pypdf import PdfReader


def read_pdf_page_count(path: Path) -> int:
    return len(PdfReader(str(path)).pages)


def read_docx_word_count(path: Path) -> int:
    """Includes table cell text, not just paragraph text -- Gate 1's parity
    check (Task 11) compares this against len(output_body.split()), and the
    OUTPUT markdown's word count includes every pipe/cell token in a
    rendered table. Counting only paragraphs here would systematically
    undercount any docx with a substantial table, producing a false
    word_count_parity_failed rejection on a perfectly good conversion."""
    document = docx.Document(str(path))
    word_count = sum(len(paragraph.text.split()) for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                word_count += len(cell.text.split())
    return word_count


def read_docx_table_count(path: Path) -> int:
    return len(docx.Document(str(path)).tables)


def read_xlsx_sheet_and_row_counts(path: Path) -> tuple[int, int]:
    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    try:
        sheet_count = len(workbook.sheetnames)
        row_count_total = 0
        for name in workbook.sheetnames:
            sheet = workbook[name]
            row_count_total += sum(
                1 for row in sheet.iter_rows() if any(cell.value is not None for cell in row)
            )
        return sheet_count, row_count_total
    finally:
        workbook.close()
