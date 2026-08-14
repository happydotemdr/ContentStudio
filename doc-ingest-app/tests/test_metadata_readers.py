import docx
import openpyxl
from pypdf import PdfWriter

from doc_ingest import metadata_readers


def test_read_pdf_page_count(tmp_path):
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_blank_page(width=200, height=200)
    writer.add_blank_page(width=200, height=200)
    pdf_path = tmp_path / "three_pages.pdf"
    with open(pdf_path, "wb") as fh:
        writer.write(fh)
    assert metadata_readers.read_pdf_page_count(pdf_path) == 3


def test_read_docx_word_count(tmp_path):
    document = docx.Document()
    document.add_paragraph("one two three four five")
    document.add_paragraph("six seven")
    docx_path = tmp_path / "sample.docx"
    document.save(docx_path)
    assert metadata_readers.read_docx_word_count(docx_path) == 7


def test_read_docx_word_count_includes_table_cell_text(tmp_path):
    document = docx.Document()
    document.add_paragraph("one two three")  # 3 paragraph words
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "four five"    # 2 table words
    table.rows[0].cells[1].text = "six seven eight"  # 3 table words
    docx_path = tmp_path / "sample.docx"
    document.save(docx_path)
    assert metadata_readers.read_docx_word_count(docx_path) == 8  # 3 + 2 + 3


def test_read_docx_table_count(tmp_path):
    document = docx.Document()
    document.add_paragraph("intro")
    document.add_table(rows=2, cols=2)
    document.add_table(rows=1, cols=3)
    docx_path = tmp_path / "with_tables.docx"
    document.save(docx_path)
    assert metadata_readers.read_docx_table_count(docx_path) == 2


def test_read_xlsx_sheet_and_row_counts(tmp_path):
    workbook = openpyxl.Workbook()
    sheet1 = workbook.active
    sheet1.title = "Sheet1"
    for i in range(5):
        sheet1.append([i, i * 2])
    sheet2 = workbook.create_sheet("Sheet2")
    for i in range(3):
        sheet2.append([i])
    xlsx_path = tmp_path / "sample.xlsx"
    workbook.save(xlsx_path)
    sheet_count, row_count = metadata_readers.read_xlsx_sheet_and_row_counts(xlsx_path)
    assert sheet_count == 2
    assert row_count == 8
